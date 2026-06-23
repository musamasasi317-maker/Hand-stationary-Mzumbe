import streamlit as st
import requests
import io
import re
import os
import base64
import copy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.oxml.text.paragraph import CT_P
from docx.opc.packuri import PackURI
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# -------------------- Page Configuration --------------------
st.set_page_config(
    page_title="FormatFix Premium",
    page_icon="📄",
    layout="centered",
    initial_sidebar_state="collapsed"
)

# -------------------- Custom CSS Styling --------------------
st.markdown("""
<style>
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    h1, h2, h3 {
        color: #1a2a6c;
        font-weight: 700;
    }
    .stButton button {
        background-color: #1a2a6c;
        color: white;
        font-weight: bold;
        border-radius: 10px;
        padding: 0.5rem 2rem;
        border: none;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: all 0.3s ease;
    }
    .stButton button:hover {
        transform: scale(1.02);
        background-color: #0f1a4a;
        box-shadow: 0 6px 10px rgba(0,0,0,0.2);
    }
    .stRadio label {
        font-weight: 600;
        color: #1a2a6c;
    }
    .stTextInput input, .stTextArea textarea {
        border-radius: 8px;
        border: 1px solid #ced4da;
        padding: 10px;
    }
    .stSuccess {
        background-color: #d4edda;
        border-radius: 10px;
        padding: 10px;
    }
    .footer {
        color: #6c757d;
        text-align: center;
        margin-top: 30px;
    }
    .preview-frame {
        width: 100%;
        height: 600px;
        border: 2px solid #1a2a6c;
        border-radius: 10px;
        box-shadow: 0 8px 16px rgba(0,0,0,0.2);
    }
</style>
""", unsafe_allow_html=True)

# -------------------- XML Helpers --------------------
def add_mzumbe_double_border_to_first_section(doc):
    if len(doc.sections) == 0:
        return
    section = doc.sections[0]
    sect_pr = section._sectPr
    pg_borders = OxmlElement('w:pgBorders')
    pg_borders.set(qn('w:display'), 'firstPage')
    pg_borders.set(qn('w:offsetFrom'), 'text')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'double')
        border.set(qn('w:sz'), '24')
        border.set(qn('w:space'), '4')
        border.set(qn('w:color'), 'auto')
        pg_borders.append(border)
    sect_pr.append(pg_borders)

def remove_borders_from_other_sections(doc):
    for idx, section in enumerate(doc.sections):
        if idx == 0:
            continue
        sect_pr = section._sectPr
        for elem in sect_pr.findall(qn('w:pgBorders')):
            sect_pr.remove(elem)

def add_page_numbers(doc):
    for idx, section in enumerate(doc.sections):
        if idx == 0:
            section.different_first_page_header_footer = True
            first_footer = section.first_page_footer
            first_footer.paragraphs.clear()
            footer = section.footer
            footer.paragraphs.clear()
        else:
            footer = section.footer
            footer.paragraphs.clear()
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            fld_simple = OxmlElement('w:fldSimple')
            fld_simple.set(qn('w:instr'), 'PAGE')
            run._r.append(fld_simple)
            sect_pr = section._sectPr
            pg_num_type = OxmlElement('w:pgNumType')
            pg_num_type.set(qn('w:start'), '1')
            sect_pr.append(pg_num_type)

# -------------------- Logo Loader --------------------
def get_logo_stream(uploaded_logo=None):
    if uploaded_logo is not None:
        return io.BytesIO(uploaded_logo.getvalue())
    local_file = "mzumbe_logo.png"
    if os.path.exists(local_file):
        with open(local_file, 'rb') as f:
            return io.BytesIO(f.read())
    urls = [
        "https://upload.wikimedia.org/wikipedia/commons/e/e1/Mzumbe_University_logo.png",
        "https://www.mzumbe.ac.tz/images/logo.png",
    ]
    for url in urls:
        try:
            response = requests.get(url, timeout=10)
            if response.status_code == 200:
                return io.BytesIO(response.content)
        except Exception:
            continue
    return None

# -------------------- Reference Formatter --------------------
def format_reference_paragraph(paragraph):
    text = paragraph.text
    if not text.strip():
        return
    import re
    match_year = re.search(r'\((\d{4})\)\.\s*', text)
    if not match_year:
        return
    start_idx = match_year.end()
    remainder = text[start_idx:]
    match_pub = re.search(r'\.\s+([A-Z][a-z])', remainder)
    if match_pub:
        pub_start = start_idx + match_pub.start() + 1
        title = text[start_idx:pub_start].strip()
        publisher = text[pub_start:].strip()
        before_title = text[:start_idx].strip()
        paragraph.clear()
        if before_title:
            run = paragraph.add_run(before_title)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        run = paragraph.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        if publisher:
            if not publisher.startswith('.'):
                run = paragraph.add_run('. ')
            else:
                run = paragraph.add_run(' ')
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run = paragraph.add_run(publisher)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    else:
        before_title = text[:start_idx].strip()
        title = remainder.strip()
        paragraph.clear()
        if before_title:
            run = paragraph.add_run(before_title)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
        run = paragraph.add_run(title)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True

# -------------------- Cover Page Builder --------------------
def create_cover_page(doc, assignment_type, metadata, students=None, uploaded_logo=None, case_mode="original"):
    def apply_case(text):
        if case_mode == "capital":
            return text.upper()
        elif case_mode == "lower":
            return text.lower()
        elif case_mode == "sentence":
            if not text:
                return text
            text_lower = text.lower()
            import re
            return re.sub(r'(^|\.\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text_lower)
        return text

    def add_centered(text, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(apply_case(text))
        run.bold = bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    add_centered("MZUMBE UNIVERSITY", bold=True)

    logo_stream = get_logo_stream(uploaded_logo)
    if logo_stream:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_stream, width=Inches(1.8))
    else:
        p = add_centered("[Nembo ya Mzumbe - tafadhali weka picha mwenyewe]", bold=False)
        if not hasattr(st, '_logo_hint_shown'):
            st.info("💡 Kidokezo: Weka faili 'mzumbe_logo.png' kwenye folder ya app, au pakia picha kupitia chaguo hapa chini.")
            st._logo_hint_shown = True

    add_centered("SCHOOL OF BUSINESS (SOB)", bold=True)

    for _ in range(2):
        doc.add_paragraph()

    # Build rows
    if assignment_type == "Individual Assignment":
        rows_data = [
            ("PROGRAMME :", metadata.get("programme", "")),
            ("SUBJECT NAME :", metadata.get("subject_name", "")),
            ("SUBJECT CODE :", metadata.get("subject_code", "")),
            ("LECTURER'S NAME :", metadata.get("lecturer", "")),
            ("NAME :", metadata.get("student_name", "")),
            ("REG.NUMBER :", metadata.get("reg_number", "")),
            ("NATURE OF WORK :", "INDIVIDUAL ASSIGNMENT"),
            ("DATE OF SUBMISSION :", metadata.get("date_submission", "")),
        ]
    else:
        rows_data = [
            ("PROGRAMME :", metadata.get("programme", "")),
            ("GROUP NUMBER :", metadata.get("group_number", "")),
            ("SUBJECT NAME :", metadata.get("subject_name", "")),
            ("SUBJECT CODE :", metadata.get("subject_code", "")),
            ("LECTURER'S NAME :", metadata.get("lecturer", "")),
            ("NATURE OF TASK :", "GROUP ASSIGNMENT"),
            ("DATE OF SUBMISSION :", metadata.get("date_submission", "")),
        ]

    # Create borderless table
    table = doc.add_table(rows=len(rows_data), cols=2)
    table.autofit = False
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        # Label
        cell_label = row.cells[0]
        p_label = cell_label.paragraphs[0]
        p_label.paragraph_format.space_after = Pt(6)
        p_label.paragraph_format.line_spacing = 2.5
        run_l = p_label.add_run(apply_case(label))
        run_l.bold = True
        run_l.font.name = 'Times New Roman'
        run_l.font.size = Pt(12)

        # Value
        cell_value = row.cells[1]
        p_value = cell_value.paragraphs[0]
        p_value.paragraph_format.space_after = Pt(6)
        p_value.paragraph_format.line_spacing = 2.5
        run_v = p_value.add_run(apply_case(str(value)))
        run_v.font.name = 'Times New Roman'
        run_v.font.size = Pt(12)

    # Remove all borders (invisible table)
    tblPr = table._tbl.tblPr
    borders_xml = parse_xml(f'<w:tblBorders {nsdecls("w")}>'
                            f'<w:top w:val="none"/>'
                            f'<w:left w:val="none"/>'
                            f'<w:bottom w:val="none"/>'
                            f'<w:right w:val="none"/>'
                            f'<w:insideH w:val="none"/>'
                            f'<w:insideV w:val="none"/>'
                            f'</w:tblBorders>')
    tblPr.append(borders_xml)

    # Student table for Group Assignment (visible)
    if assignment_type == "Group Assignment" and students:
        doc.add_paragraph()
        student_table = doc.add_table(rows=1, cols=3)
        student_table.style = 'Table Grid'
        hdr_cells = student_table.rows[0].cells
        hdr_cells[0].text = "S/N"
        hdr_cells[1].text = "NAME OF STUDENTS"
        hdr_cells[2].text = "REGISTRATION NUMBER"
        for idx, (name, reg) in enumerate(students, start=1):
            row_cells = student_table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = apply_case(name.strip())
            row_cells[2].text = apply_case(reg.strip())
        for row in student_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                    if not p.runs:
                        run = p.add_run(p.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    for _ in range(3):
        doc.add_paragraph()

# -------------------- Clone Content with Images/Shapes --------------------
def clone_body_content(original_doc, new_doc, case_mode="original"):
    """
    Clone all body elements (paragraphs, tables) from original_doc to new_doc,
    preserving images, shapes, and all formatting.
    Also applies case transformation to text runs.
    """
    def apply_case_text(text):
        if case_mode == "capital":
            return text.upper()
        elif case_mode == "lower":
            return text.lower()
        elif case_mode == "sentence":
            if not text:
                return text
            import re
            text_lower = text.lower()
            return re.sub(r'(^|\.\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text_lower)
        return text

    # Get the body element of the new document
    new_body = new_doc._element.body

    # Iterate over original body children
    for child in original_doc._element.body.iterchildren():
        if child.tag == qn('w:p'):  # Paragraph
            # Create a deep copy of the paragraph element
            new_child = copy.deepcopy(child)
            # Apply case transformation to text runs
            for run in new_child.findall('.//' + qn('w:r')):
                t_elem = run.find(qn('w:t'))
                if t_elem is not None and t_elem.text:
                    t_elem.text = apply_case_text(t_elem.text)
            # Append to new body
            new_body.append(new_child)
        elif child.tag == qn('w:tbl'):  # Table
            new_child = copy.deepcopy(child)
            # Apply case to all text in table cells
            for cell in new_child.findall('.//' + qn('w:tc')):
                for p in cell.findall(qn('w:p')):
                    for run in p.findall(qn('w:r')):
                        t_elem = run.find(qn('w:t'))
                        if t_elem is not None and t_elem.text:
                            t_elem.text = apply_case_text(t_elem.text)
            new_body.append(new_child)

    # Now we need to copy the relationship parts (images) from original to new document
    # We'll copy all image parts and update relationship IDs.
    # This is complex; we'll iterate over all relationships in original part and add them to new part.
    # We'll also need to update the rId references in the cloned elements.
    # We'll do this by iterating over all 'blip' elements and replacing their rId with a new one.
    # Get the new document's part (the main document part)
    new_part = new_doc.part
    old_part = original_doc.part

    # Map old rId to new rId
    rId_map = {}

    # Iterate over all blips in the new document (they have old rIds)
    for blip in new_doc._element.findall('.//' + qn('a:blip')):
        old_rId = blip.get(qn('r:embed'))
        if old_rId is None:
            continue
        if old_rId in rId_map:
            new_rId = rId_map[old_rId]
        else:
            # Get the image part from the original document
            old_image_part = old_part.related_parts[old_rId]
            # Add this image part to the new document with a new relationship
            # We need to create a new part of the same content type
            # Get the blob
            blob = old_image_part.blob
            # Determine content type (image/jpeg, image/png, etc.) from part
            content_type = old_image_part.content_type
            # Create a new part in the new document
            # Generate a unique part name
            partname = f"/word/media/image{len(rId_map)+1}.png"
            # But we need to know the correct extension; we'll just use the same part name as original? We can use a new unique name.
            # We'll use a counter.
            # We'll add a part to the package
            new_part.package.part.related_parts  # Not straightforward.
            # Simpler: we can reuse the same blob and add a new relationship.
            # We'll use the package's get_or_add_image_part method? Not public.
            # We'll use a helper: add a new image part.
            # We'll use the existing image part's content type and blob.
            # We'll create a new Part and add it.
            from docx.opc.part import Part
            from docx.opc.packuri import PackURI
            new_uri = PackURI(f"/word/media/image{len(rId_map)+1}.img")
            new_part_obj = Part(new_uri, content_type, blob)
            # Add this part to the package
            new_part.package.parts.append(new_part_obj)
            # Create relationship between new_doc part and this new part
            new_rId = new_part.relate_to(new_part_obj, RT.IMAGE)
            rId_map[old_rId] = new_rId
        # Update the blip's r:embed attribute with the new rId
        blip.set(qn('r:embed'), new_rId)

# -------------------- Core Processing --------------------
def process_document(input_mode, uploaded_file, raw_text, metadata, assignment_type, students, uploaded_logo, case_mode):
    new_doc = Document()

    # Cover page
    create_cover_page(new_doc, assignment_type, metadata, students, uploaded_logo, case_mode)

    # Add section break for content
    new_doc.add_section()

    if input_mode == "Pakia Faili la DOCX" and uploaded_file is not None:
        original = Document(uploaded_file)
        # Clone all body content (preserving shapes, images, tables)
        clone_body_content(original, new_doc, case_mode)
    elif input_mode == "Bandika Maandishi" and raw_text:
        # For pasted text, we create paragraphs manually (no images)
        lines = [line.strip() for line in raw_text.splitlines() if line.strip() != ""]
        in_reference = False
        for text in lines:
            if text.upper() in ["REFERENCES", "BIBLIOGRAPHY"]:
                new_doc.add_page_break()
                in_reference = True
                p_heading = new_doc.add_paragraph()
                p_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p_heading.add_run(apply_case_text(text, case_mode))
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                continue
            new_para = new_doc.add_paragraph()
            run = new_para.add_run(apply_case_text(text, case_mode))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            new_para.paragraph_format.line_spacing = 1.5
            if in_reference and text.upper() not in ["REFERENCES", "BIBLIOGRAPHY"]:
                new_para.paragraph_format.first_line_indent = Inches(-0.5)
                new_para.paragraph_format.left_indent = Inches(0.5)
                format_reference_paragraph(new_para)
    else:
        raise ValueError("Hakuna maandishi yaliyopatikana.")

    # Apply page borders (first page only) and page numbers
    add_mzumbe_double_border_to_first_section(new_doc)
    remove_borders_from_other_sections(new_doc)
    add_page_numbers(new_doc)

    buffer = io.BytesIO()
    new_doc.save(buffer)
    buffer.seek(0)
    return buffer

def apply_case_text(text, mode):
    if mode == "capital":
        return text.upper()
    elif mode == "lower":
        return text.lower()
    elif mode == "sentence":
        if not text:
            return text
        import re
        text_lower = text.lower()
        return re.sub(r'(^|\.\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text_lower)
    return text

# -------------------- PDF Preview (optional) --------------------
def convert_docx_to_pdf(docx_buffer):
    try:
        from docx2pdf import convert
        import tempfile
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(docx_buffer.getvalue())
            tmp_docx_path = tmp_docx.name
        pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        convert(tmp_docx_path, pdf_path)
        with open(pdf_path, "rb") as f:
            pdf_buffer = io.BytesIO(f.read())
        os.unlink(tmp_docx_path)
        os.unlink(pdf_path)
        pdf_buffer.seek(0)
        return pdf_buffer
    except ImportError:
        raise ImportError("Moduli ya docx2pdf haijasakinishwa. Endesha: pip install docx2pdf")
    except Exception as e:
        raise Exception(f"Hitilafu: {e}")

# -------------------- Streamlit UI --------------------
st.markdown(
    """
    <h1 style='text-align: center;'>📄 FormatFix Premium</h1>
    <p style='text-align: center; font-size: 1.2rem;'>Panga kazi za wanafunzi kwa mtindo wa kitaaluma <strong>katika sekunde moja</strong>.</p>
    <hr>
    """,
    unsafe_allow_html=True
)

# Sidebar
st.sidebar.header("⚙️ Mipangilio")
uploaded_logo = st.sidebar.file_uploader(
    "Pakia Nembo Mwenyewe (hiari)",
    type=["png", "jpg", "jpeg"],
    help="Ikiwa nembo ya mtandao haipatikani, pakia picha yako."
)

case_mode = st.sidebar.selectbox(
    "Uandishi wa Herufi",
    ["Uandishi halisi (Original)", "Herufi kubwa (CAPITAL)", "Herufi ndogo (small)", "Herufi za kwanza kubwa (Sentence)"],
    index=0
)
case_map = {
    "Uandishi halisi (Original)": "original",
    "Herufi kubwa (CAPITAL)": "capital",
    "Herufi ndogo (small)": "lower",
    "Herufi za kwanza kubwa (Sentence)": "sentence"
}
case_mode_internal = case_map[case_mode]

# Main input
input_mode = st.radio(
    "Chagua njia ya kuingiza maandishi",
    ["Pakia Faili la DOCX", "Bandika Maandishi"],
    index=0,
    horizontal=True
)

uploaded_file = None
raw_text = ""
if input_mode == "Pakia Faili la DOCX":
    uploaded_file = st.file_uploader(
        "Pakua faili la assignment ambalo halijapangwa (.docx)",
        type=["docx"],
        help="Faili la DOCX pekee linakubalika."
    )
else:
    raw_text = st.text_area(
        "Bandika maandishi ya assignment hapa",
        height=300,
        placeholder="Weka kila aya kwenye mstari mpya."
    )

assignment_type = st.radio(
    "Aina ya Kazi",
    ["Individual Assignment", "Group Assignment"],
    index=0,
    horizontal=True
)

col1, col2 = st.columns(2)
with col1:
    programme = st.text_input("Programme", placeholder="mfano: BBA-EIM 1A")
    subject_name = st.text_input("Jina la Somo", placeholder="mfano: PRINCIPLES OF MANAGEMENT")
    subject_code = st.text_input("Msimbo wa Somo", placeholder="mfano: PUB 111")
    lecturer = st.text_input("Jina la Mhadhiri", placeholder="PAULO MARO")
with col2:
    if assignment_type == "Individual Assignment":
        student_name = st.text_input("Jina la Mwanafunzi", placeholder="Jina kamili")
        reg_number = st.text_input("Namba ya Usajili", placeholder="mfano: 1739110/T.25")
        group_number = ""
    else:
        group_number = st.text_input("Namba ya Kikundi", placeholder="mfano: Group 5")
        student_name = ""
        reg_number = ""
    date_submission = st.text_input("Tarehe ya Uwasilishaji", placeholder="mfano: 15 JUNE 2026")

students = []
if assignment_type == "Group Assignment":
    st.subheader("Taarifa za Wanafunzi")
    st.caption("Weka kila mwanafunzi kwenye mstari mpya: `Jina, Namba ya Usajili` (tumia comma au slash)")
    student_text = st.text_area(
        "Wanafunzi",
        height=150,
        placeholder="John Mushi, 2024-01-0002\nJane Kilima, 2024-01-0003"
    )
    if student_text.strip():
        for line in student_text.strip().splitlines():
            line = line.strip()
            if not line:
                continue
            parts = re.split(r'[,/]\s*', line, maxsplit=1)
            if len(parts) == 2:
                students.append((parts[0].strip(), parts[1].strip()))
            else:
                st.warning(f"Imepitishwa mstari (inatarajiwa 'Jina, RegNo'): {line}")

col_btn1, col_btn2 = st.columns(2)
process_btn = col_btn1.button("🚀 Pangwa Katika Sekunde", use_container_width=True)
preview_btn = col_btn2.button("👁️ Onyesha Hati (Preview)", use_container_width=True, disabled=not st.session_state.get('processed', False))

if 'doc_buffer' not in st.session_state:
    st.session_state.doc_buffer = None
if 'processed' not in st.session_state:
    st.session_state.processed = False
if 'pdf_buffer' not in st.session_state:
    st.session_state.pdf_buffer = None

if process_btn:
    errors = []
    if input_mode == "Pakia Faili la DOCX" and uploaded_file is None:
        errors.append("Tafadhali pakua faili la DOCX.")
    elif input_mode == "Bandika Maandishi" and not raw_text.strip():
        errors.append("Tafadhali bandika maandishi.")
    if not programme:
        errors.append("Programme inahitajika.")
    if not subject_name:
        errors.append("Jina la Somo linahitajika.")
    if not subject_code:
        errors.append("Msimbo wa Somo unahitajika.")
    if not lecturer:
        errors.append("Jina la Mhadhiri linahitajika.")
    if not date_submission:
        errors.append("Tarehe ya Uwasilishaji inahitajika.")
    if assignment_type == "Individual Assignment":
        if not student_name:
            errors.append("Jina la Mwanafunzi linahitajika.")
        if not reg_number:
            errors.append("Namba ya Usajili inahitajika.")
    else:
        if not group_number:
            errors.append("Namba ya Kikundi inahitajika.")
        if not students:
            st.warning("Hakuna data ya wanafunzi. Jedwali litakuwa tupu.")

    if errors:
        for err in errors:
            st.error(err)
    else:
        try:
            with st.spinner("Inapanga hati yako... Tafadhali subiri."):
                metadata = {
                    "programme": programme,
                    "subject_name": subject_name,
                    "subject_code": subject_code,
                    "lecturer": lecturer,
                    "date_submission": date_submission,
                    "student_name": student_name,
                    "reg_number": reg_number,
                    "group_number": group_number,
                }
                buffer = process_document(
                    input_mode,
                    uploaded_file,
                    raw_text,
                    metadata,
                    assignment_type,
                    students,
                    uploaded_logo,
                    case_mode_internal
                )
                st.session_state.doc_buffer = buffer
                st.session_state.processed = True
                st.success("🎉 Kazi imepangwa kikamilifu!")
                st.balloons()

                st.download_button(
                    label="⬇ Pakua Hati Iliyopangwa (DOCX)",
                    data=buffer,
                    file_name="FormatFix_Perfect_Document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

                # Preview
                try:
                    with st.spinner("Inaandaa onyesho la PDF..."):
                        pdf_buffer = convert_docx_to_pdf(buffer)
                        st.session_state.pdf_buffer = pdf_buffer
                        base64_pdf = base64.b64encode(pdf_buffer.getvalue()).decode('utf-8')
                        pdf_display = f'<iframe class="preview-frame" src="data:application/pdf;base64,{base64_pdf}" type="application/pdf"></iframe>'
                        st.markdown(pdf_display, unsafe_allow_html=True)
                except ImportError:
                    st.info("💡 Ili kuona preview, hakikisha una Word na moduli ya docx2pdf. Endesha: `pip install docx2pdf`")
                except Exception as e:
                    st.warning(f"Haikuweza kuonyesha preview: {e}")

        except Exception as e:
            st.error(f"Hitilafu ilitokea: {e}")
            st.exception(e)

if preview_btn and st.session_state.processed and st.session_state.doc_buffer is not None:
    try:
        with st.spinner("Inaandaa onyesho la PDF..."):
            pdf_buffer = convert_docx_to_pdf(st.session_state.doc_buffer)
            st.session_state.pdf_buffer = pdf_buffer
            base64_pdf = base64.b64encode(pdf_buffer.getvalue()).decode('utf-8')
            pdf_display = f'<iframe class="preview-frame" src="data:application/pdf;base64,{base64_pdf}" type="application/pdf"></iframe>'
            st.markdown(pdf_display, unsafe_allow_html=True)
    except ImportError:
        st.error("Moduli ya docx2pdf haijasakinishwa. Endesha: pip install docx2pdf")
    except Exception as e:
        st.error(f"Haikuweza kuonyesha preview: {e}")

st.markdown(
    """
    <hr>
    <p class="footer">Imetengenezwa kwa ❤️ na Musa Lucas Masasi · Chuo Kikuu cha Mzumbe</p>
    """,
    unsafe_allow_html=True
)
